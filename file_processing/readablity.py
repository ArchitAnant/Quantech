import PyPDF2
import os
import fitz
from docx import Document


def read_doc(file_path):
    parra_lines = []
    # Open the .docx file
    doc = Document(file_path)
    # Extract and print text
    for paragraph in doc.paragraphs:
        parra_lines.append(paragraph.text)

    return " ".join(parra_lines)

def read_pdf(file_path):
    lines = []

# Open the PDF file
    with open(file_path, 'rb') as file:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(file)

        # Get the number of pages
        num_pages = len(pdf_reader.pages)
        print(f"Number of pages: {num_pages}")

        # Extract text from each page
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            lines.append(text)

    return " ".join(lines)


def check_images_in_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    has_images = False

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        image_list = page.get_images(full=True)
        if image_list:
            has_images = True
            break

    return has_images


def check_docx_for_images(docx_path):
    doc = Document(docx_path)
    has_images = False

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            has_images = True
            break

    return has_images



def get_file_size(file_path):
    try:
        return os.path.getsize(file_path)
    except OSError as e:
        print(f"Error: {e}")
        return None
    
def images_in_pdf(pdf_file):
    # return images_between_text, image_dimensions of the images at the end and beggining of every page of the pdf
    pdf_document = fitz.open(pdf_file)
    images_between_text = False
    image_dimensions = {}

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        image_list = page.get_images(full=True)
        if image_list:
            images_between_text = True
            break
        else:
            image_dimensions[f'page_{page_number}_start_image'] = page.get_image_list()[0][2]
            image_dimensions[f'page_{page_number}_end_image'] = page.get_image_list()[-1][2]

    
    return images_between_text, image_dimensions


def images_in_docx(docx_file):
    #return images_between_text, image_dimensions of the images at the end and beggining of every page of the docx
    doc = Document(docx_file)
    images_between_text = False
    image_dimensions = {}

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            images_between_text = True
            break
        else:
            image_dimensions[f'page_{i}_start_image'] = get_image_dimensions(rel.target)
            image_dimensions[f'page_{i}_end_image'] = get_image_dimensions(rel.target)
    
    return images_between_text, image_dimensions
    

    

def get_image_dimensions(element):
    for el in element.iter():
        if el.tag.endswith('extents'):
            # Convert EMUs to points
            cx = int(el.attrib['cx']) * (72 / 914400)  # Convert EMUs to points
            cy = int(el.attrib['cy']) * (72 / 914400)
            return cx, cy


def is_readable(file_path):
    pdf = False
    if ".pdf" in file_path:
        str_val = read_pdf(file_path)
        pdf = True
    else:
        str_val = read_doc(file_path)
        print(len(str_val))
    
    if len(str_val)==0:
        return False

    else:
        if pdf and (not check_images_in_pdf(file_path)):
            return True
        elif (not pdf) and (not check_docx_for_images(file_path)):
            return True
        
        # temp
        else :
            return False
        
        #Multiple images in between text implementation pending

        # elif pdf:
        #     image_in_between,dimensions = images_in_pdf(file_path) 
        #     print(dimensions)
        #     if image_in_between:
        #         return False
        #     else:
        #         for image in dimensions.keys():
        #             if dimensions[image][1] > 160:
        #                 return False
        # else:
        #     image_in_between,dimensions = images_in_docx(file_path) 
        #     print(dimensions)
        #     if image_in_between:
        #         return False
        #     else:
        #         for image in dimensions.keys():
        #             if dimensions[image][1] > 160:
        #                 return False



# file_path = "/Users/architanant/Documents/Quantech/file_processing/test.docx"
# print(is_readable(file_path))
    


